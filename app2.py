import io
import pandas as pd
import streamlit as st
from langchain.text_splitter import RecursiveCharacterTextSplitter
from langchain.vectorstores import FAISS
from langchain.chains.question_answering import load_qa_chain
from langchain.prompts import PromptTemplate
from langchain_community.document_loaders.csv_loader import CSVLoader
from langchain_google_genai import ChatGoogleGenerativeAI
from langchain_google_genai import GoogleGenerativeAIEmbeddings
from dotenv import load_dotenv
import os
from docx import Document
from datetime import datetime

# Load environment variables
load_dotenv()
GOOGLE_API_KEY = os.getenv("GOOGLE_API_KEY")

# Initialize the LLM using Google Gemini
llm = ChatGoogleGenerativeAI(
    model="gemini-pro",
    google_api_key=GOOGLE_API_KEY,
    temperature=0
)

def create_word_document(analysis_text):
    doc = Document()
    
    # Add title
    doc.add_heading('Sales Analysis Report', 0)
    
    # Add timestamp
    timestamp = datetime.now().strftime("%Y-%m-%d %H:%M:%S")
    doc.add_paragraph(f'Generated on: {timestamp}')
    
    # Add a line break
    doc.add_paragraph()
    
    # Split the analysis text into sections based on numbered points
    sections = analysis_text.split('\n\n')
    
    for section in sections:
        # Check if the section is a main heading (starts with a number followed by a dot)
        if any(section.startswith(f"{i}.") for i in range(1, 6)):
            doc.add_heading(section.split('\n')[0], level=1)
            # Add the rest of the section content
            content = '\n'.join(section.split('\n')[1:])
            if content.strip():
                doc.add_paragraph(content.strip())
        else:
            # Add regular paragraphs
            if section.strip():
                doc.add_paragraph(section.strip())
    
    # Save to bytes buffer
    doc_buffer = io.BytesIO()
    doc.save(doc_buffer)
    doc_buffer.seek(0)
    return doc_buffer

def get_csv_text(csv_docs):
    # Reading the uploaded CSV file and returning its content and dataframe
    all_text = ""
    dfs = []
    for uploaded_file in csv_docs:
        df = pd.read_csv(uploaded_file)
        dfs.append(df)
        all_text += df.to_string()
    return all_text, dfs

def get_text_chunks(text):
    text_splitter = RecursiveCharacterTextSplitter(chunk_size=10000, chunk_overlap=1000)
    chunks = text_splitter.split_text(text)
    return chunks

def get_vector_store(text_chunks):
    embeddings = GoogleGenerativeAIEmbeddings(model="models/embedding-001")
    vector_store = FAISS.from_texts(text_chunks, embedding=embeddings)
    vector_store.save_local("faiss_index")

def generate_initial_analysis(df):
    analysis_prompt = """Generate a comprehensive sales analysis report based on the provided data. 
    The report should include:

    1. Overall Performance Metrics:
       - Total revenue
       - Total number of transactions
       - Average transaction value
       - Revenue trends over time

    2. Product Analysis:
       - Top performing products by revenue
       - Top performing products by quantity
       - Product category performance (if applicable)

    3. Sales Patterns:
       - Monthly/Seasonal trends
       - Peak sales periods
       - Growth rate analysis

    4. Customer Insights (if available):
       - Customer segments
       - Purchase frequency
       - Customer lifetime value

    5. Recommendations:
       - Key areas for improvement
       - Growth opportunities
       - Risk factors to monitor

    Data Context: {context}
    
    Please provide the analysis in a clear, structured format with specific numbers and percentages where relevant.
    """
    
    context = df.describe().to_string() + "\n\nColumns available: " + ", ".join(df.columns)
    
    # Create a temporary chain for initial analysis
    prompt = PromptTemplate(template=analysis_prompt, input_variables=["context"])
    chain = load_qa_chain(llm, chain_type="stuff", prompt=prompt)
    
    # Create a dummy document with the context
    from langchain.docstore.document import Document
    doc = Document(page_content=context)
    
    response = chain(
        {"input_documents": [doc], "question": "Generate initial analysis"},
        return_only_outputs=True
    )
    
    return response.get("output_text", "No analysis generated.")

def get_conversational_chain():
    prompt_template = """You are a professional sales analyst. Based on the provided sales data, answer the following question
    or generate analysis as requested. Consider all relevant metrics and provide specific numbers, percentages, and trends 
    where applicable.

    Context about the data: {context}
    Question: {question}

    Please provide a detailed, professional response that includes:
    1. Direct answer to the question
    2. Supporting data and metrics
    3. Relevant trends or patterns
    4. Business implications when applicable
    """

    prompt = PromptTemplate(template=prompt_template, input_variables=["context", "question"])
    chain = load_qa_chain(llm, chain_type="stuff", prompt=prompt)
    return chain

def user_input(user_question):
    embeddings = GoogleGenerativeAIEmbeddings(model="models/embedding-001")
    new_db = FAISS.load_local("faiss_index", embeddings, allow_dangerous_deserialization=True)
    docs = new_db.similarity_search(user_question)

    chain = get_conversational_chain()
    response = chain(
        {"input_documents": docs, "question": user_question},
        return_only_outputs=True
    )

    st.write("Reply: ", response.get("output_text", "No response generated."))

def main():
    st.set_page_config(page_title="Automated Sales Analysis Report 📊", layout="wide")
    st.header("Automated Sales Analysis Report 📊")

    with st.sidebar:
        st.title("Upload Data")
        csv_docs = st.file_uploader(
            "Upload your CSV Files and Click on the Submit & Process Button",
            accept_multiple_files=True,
            type=["csv"]
        )
        if st.button("Submit & Process"):
            if not csv_docs:
                st.warning("Please upload at least one CSV file.")
                return
            with st.spinner("Processing..."):
                raw_text, dfs = get_csv_text(csv_docs)
                if not raw_text:
                    st.error("No text found in the uploaded CSV files.")
                    return
                text_chunks = get_text_chunks(raw_text)
                get_vector_store(text_chunks)
                
                # Generate and display initial analysis
                st.session_state['initial_analysis'] = generate_initial_analysis(dfs[0])
                st.success("Processing completed successfully!")

    # Main content area
    if 'initial_analysis' in st.session_state:
        st.subheader("Automated Sales Analysis Report")
        st.write(st.session_state['initial_analysis'])
        
        # Add download button for Word document
        doc_buffer = create_word_document(st.session_state['initial_analysis'])
        st.download_button(
            label="Download Analysis as Word Document",
            data=doc_buffer,
            file_name="sales_analysis_report.docx",
            mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
        )
        
        st.divider()

if __name__ == "__main__":
    main()